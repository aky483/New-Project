import streamlit as st
import google.generativeai as genai
import os
import PyPDF2 as pdf
import requests
from bs4 import BeautifulSoup
import json
from docx import Document
from io import BytesIO
import re
import plotly.graph_objects as go

# Use secret from Streamlit Cloud
genai.configure(
    api_key=st.secrets["gemini"]["api_key"],
    transport="rest",
    client_options={"api_endpoint": "generativelanguage.googleapis.com"}
)

# Function to generate response from Gemini
def get_gemini_response(prompt):
    model = genai.GenerativeModel('gemini-2.5-flash')
    response = model.generate_content(prompt)
    return response.text

# Extract text from PDF
def input_pdf_text(uploaded_file):
    reader = pdf.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Extract text from docx
def extract_text_from_docx(file):
    doc = Document(file)
    return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

# Prompt template
prompt_template = """
You are an advanced Applicant Tracking System (ATS) simulator used by Fortune 500 companies.
Evaluate how well the candidate's resume aligns with the provided job description. 
Your goal is to help the candidate reach a 90-95% match by identifying:
- Missing exact-match keywords or phrases from the JD
- Lack of relevant tools/tech used in the job description
- Opportunities to restructure experience to match role demands

Give output in this structured JSON:
{{
  "JD Match": "85%", 
  "MissingKeywords": ["Snowflake", "CI/CD", "Data Lake"],
  "Profile Summary": "Experienced Data Engineer with 8+ years of designing scalable pipelines..."
}}

Resume:
{text}

Job Description:
{jd}
"""

# Streamlit app
st.title("Smart ATS")
st.text("Improve Your Resume ATS Compatibility")

# JD text input
jd = st.text_area("Paste the Job Description here")
if jd.strip():
    st.subheader("📝 Job Description Preview")
    st.code(jd, language="markdown")


# Resume upload
uploaded_files = st.file_uploader(
    "Upload Resumes (PDF or Word)", 
    type=["pdf", "docx"], 
    accept_multiple_files=True, 
    help="Upload multiple resumes in PDF or DOCX format"
)

submit = st.button("Submit")

if submit:
    if uploaded_files and jd.strip():
        scores = []
        result_data = []

        with st.spinner("Analyzing all resumes..."):
            for file in uploaded_files:
                if file.name.endswith(".pdf"):
                    resume_text = input_pdf_text(file)
                elif file.name.endswith(".docx"):
                    resume_text = extract_text_from_docx(file)
                else:
                    resume_text = ""

                filled_prompt = prompt_template.format(text=resume_text, jd=jd)
                response = get_gemini_response(filled_prompt)

                # Extract score
                match = re.search(r'"JD Match"\s*:\s*"(\d+)%"', response)
                score = int(match.group(1)) if match else 0

                scores.append({
                    "name": file.name,
                    "score": score,
                    "response": response
                })
                if score < 90:
                    st.warning(f"{file.name} scored {score}%. You may want to enhance it further for better ATS match.")

        # Display results
        for data in scores:
            st.subheader(f"📄 {data['name']}")

            try:
                # Try to extract JSON-like string from the Gemini response
                json_text_match = re.search(r'\{.*\}', data['response'], re.DOTALL)
                if not json_text_match:
                    raise ValueError("No JSON object found in response")

                json_text = json_text_match.group(0)
                result_json = json.loads(json_text)

                jd_match = result_json.get("JD Match", "N/A")
                missing_keywords = result_json.get("MissingKeywords", [])

                # Display JD match score
                st.markdown(f"**✅ JD Match:** {jd_match}")
                st.progress(int(jd_match.replace('%', '')))

                # Display missing keywords as bullets
                if missing_keywords:
                    st.markdown("**❌ Missing Keywords:**")
                    for keyword in missing_keywords:
                        st.markdown(f"- {keyword}")
                else:
                    st.success("No missing keywords found. Great match! ✅")

            except Exception as e:
                st.error(f"Error parsing response: {e}")
                st.code(data['response'])


        # Compare with bar chart
        st.subheader("📊 Resume Comparison Chart")
        fig = go.Figure([go.Bar(
            x=[d['name'] for d in scores],
            y=[d['score'] for d in scores],
            marker_color='teal'
        )])
        fig.update_layout(title="JD Match Score per Resume", yaxis=dict(title="Match %"), xaxis=dict(title="Resume"))
        st.plotly_chart(fig)

    else:
        st.warning("Please upload at least one resume and provide a job description.")


#Resume Generation
resume_improvement_prompt = """
You are a professional resume writer and an expert in ATS optimization.

Your job is to:
1. Parse the given resume and extract all relevant experience.
2. Carefully read the job description to extract **all important keywords**, tools, certifications, and verbs.
3. Identify **gaps** between the resume and JD.
4. Generate a **revised resume** that:
   - Organically integrates **JD-specific keywords** throughout the resume (especially in **Work Experience** and **Key Skills**).
   - Strengthens **Work Experience** by embedding JD-relevant tools, action verbs, and achievements into each bullet point.
   - Uses **quantified impact statements** wherever possible (e.g., "Improved query performance by 40%", "Reduced pipeline failure rate by 25%").
   - Matches **job title language** and avoids generic phrasing like "worked on" or "responsible for".

**Resume Format** (keep it simple: no tables, icons, or graphics):
- Name
- Professional Summary (3-5 lines using JD language)
- Key Skills (15-20 exact phrases from the JD)
- Work Experience (Chronological, each with 4-6 bullets using JD-aligned terms, metrics, action verbs)
- Education
- Certifications (only if found in resume or JD)

⚠️ Be strict about using **specific tools, methods, platforms, and measurable impact** from the JD. Do not generalize. Do not invent projects, but you may refine existing points to better match the JD.

Resume Content:
{text}

Job Description:
{jd}
"""
st.header("🧠 Enhance My Resume Based on JD")

uploaded_resume = st.file_uploader("Upload Your Existing Resume (PDF or Word)", type=["pdf", "docx"], key="resume_improve")
jd_input = st.text_area("Paste the Job Description", key="improve_jd")

improve_btn = st.button("Generate Improved Resume")

if improve_btn:
    if uploaded_resume and jd_input.strip():
        with st.spinner("Extracting and improving your resume..."):
            # Extract resume text
            if uploaded_resume.name.endswith(".pdf"):
                reader = pdf.PdfReader(uploaded_resume)
                resume_text = ""
                for page in reader.pages:
                    resume_text += page.extract_text()
            elif uploaded_resume.name.endswith(".docx"):
                resume_text = extract_text_from_docx(uploaded_resume)
            else:
                resume_text = ""

            # Format prompt
            final_prompt = resume_improvement_prompt.format(text=resume_text, jd=jd_input)

            # Get improved resume text from Gemini
            improved_resume_text = get_gemini_response(final_prompt)

            # Clean up the text by removing markdown formatting (** etc)
            clean_text = re.sub(r"\*\*", "", improved_resume_text)
            
            # Clean the Gemini response
            clean_text = re.sub(r'^.*?(?=\n[A-Z])', '', improved_resume_text, flags=re.DOTALL)
            clean_text = re.sub(r'Key Changes and Explanations:.*', '', clean_text, flags=re.DOTALL)
            clean_text = re.sub(r'\s*\(.*?good to have.*?\)', '', clean_text, flags=re.IGNORECASE)
            clean_text = re.sub(r"\*\*", "", clean_text)
            clean_text = re.sub(r'\n{3,}', '\n\n', clean_text).strip()


            # Split into sections
            sections = re.split(r"\n(?=[A-Z][a-zA-Z ]+:)", clean_text)

            # Create a Word Document
            doc = Document()

            # Title
            doc.add_heading('Resume', level=0)

            # Add each section
            for section in sections:
                lines = section.strip().split("\n")
                if len(lines) > 1:
                    heading = lines[0].strip()
                    doc.add_heading(heading, level=1)
                    for line in lines[1:]:
                        doc.add_paragraph(line.strip())
                else:
                    # In case of single-line sections
                    doc.add_paragraph(lines[0].strip())

            # Save to in-memory buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

        st.success("✅ Enhanced Resume Generated Successfully!")

        # Provide download button
        st.download_button(
            label="📥 Download Enhanced Resume (Word)",
            data=buffer,
            file_name="enhanced_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Optionally preview raw text
        st.subheader("🔍 Preview (Raw Text)")
        st.code(improved_resume_text, language="markdown")
    else:
        st.warning("Please upload your resume and provide the job description.")

